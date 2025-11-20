#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
🔧 GÉNÉRATEUR RAPPORTS PDF INDIVIDUELS - GIRASOLE 2025
Diagnostic Photovoltaïque - Expertise indépendante depuis 2012

Génère 52 rapports PDF brandés DiagPV conformes CDC GIRASOLE
Input : JSON checklists (terrain + BE) + photos ZIP
Output : 52 rapports PDF professionnels avec photos

Auteur : DiagPV Assistant Pro
Date : 19 novembre 2025
Version : 1.0 PRODUCTION
"""

import json
import zipfile
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List
import base64
import io

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("❌ Module 'python-docx' manquant. Installation:")
    print("   pip install python-docx")
    exit(1)

try:
    from PIL import Image
except ImportError:
    print("❌ Module 'Pillow' manquant. Installation:")
    print("   pip install Pillow")
    exit(1)


# ====================================================================================================
# 🎨 CONFIGURATION IDENTITÉ DIAGPV
# ====================================================================================================

DIAGPV_CONFIG = {
    "nom": "Diagnostic Photovoltaïque",
    "groupe": "Groupe Watt&co",
    "adresse": "3 rue d'Apollo",
    "code_postal": "31240",
    "ville": "L'UNION",
    "telephone": "05.81.10.16.59",
    "email": "contact@diagpv.fr",
    "web": "www.diagnosticphotovoltaique.fr",
    "rcs": "RCS Toulouse 792 972 309",
    "assurance": "Assurance RC Professionnelle MMA",
    "baseline": "Expertise photovoltaïque indépendante depuis 2012",
    "couleur_primaire": (46, 204, 113),  # Vert #2ECC71
    "couleur_secondaire": (149, 165, 166),  # Gris #95A5A6
    "couleur_texte": (44, 62, 80)  # Bleu foncé #2C3E50
}

SIGNATAIRE = {
    "nom": "Fabien CORRERA",
    "fonction": "Responsable Technique",
    "societe": "Diagnostic Photovoltaïque"
}

DISCLAIMER = """Ce rapport a été réalisé en toute indépendance par Diagnostic Photovoltaïque.
Les conclusions techniques reflètent uniquement l'état constaté lors de l'intervention terrain.
Méthodologie conforme IEC 62446-1 | NF C 15-100 | DTU 40.35
RCS Toulouse 792 972 309 - Assurance RC Professionnelle MMA
Rapport non contractuel - Propriété intellectuelle Diagnostic Photovoltaïque"""


# ====================================================================================================
# 🔧 FONCTIONS UTILITAIRES
# ====================================================================================================

def charger_json(chemin: Path) -> Dict[str, Any]:
    """Charge un fichier JSON"""
    try:
        with open(chemin, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"❌ Erreur chargement {chemin.name}: {e}")
        return {}


def extraire_photos_zip(zip_path: Path) -> Dict[str, bytes]:
    """
    Extrait toutes les photos d'un ZIP
    Retourne dict {nom_fichier: bytes_image}
    """
    photos = {}
    try:
        with zipfile.ZipFile(zip_path, 'r') as zf:
            for filename in zf.namelist():
                if filename.lower().endswith(('.jpg', '.jpeg', '.png')):
                    photos[filename] = zf.read(filename)
        print(f"   📸 {len(photos)} photos extraites de {zip_path.name}")
        return photos
    except Exception as e:
        print(f"   ⚠️  Erreur extraction photos {zip_path.name}: {e}")
        return {}


def redimensionner_image(image_bytes: bytes, largeur_max: int = 400) -> bytes:
    """Redimensionne image pour insertion Word"""
    try:
        img = Image.open(io.BytesIO(image_bytes))
        
        # Calculer nouvelle taille
        ratio = largeur_max / img.width
        nouvelle_largeur = largeur_max
        nouvelle_hauteur = int(img.height * ratio)
        
        # Redimensionner
        img_resized = img.resize((nouvelle_largeur, nouvelle_hauteur), Image.Resampling.LANCZOS)
        
        # Convertir en bytes
        output = io.BytesIO()
        img_resized.save(output, format=img.format if img.format else 'JPEG')
        return output.getvalue()
    except Exception as e:
        print(f"   ⚠️  Erreur redimensionnement image: {e}")
        return image_bytes


# ====================================================================================================
# 📝 FONCTIONS GÉNÉRATION SECTIONS RAPPORT
# ====================================================================================================

def ajouter_entete_diagpv(doc: Document):
    """Ajoute en-tête branded DiagPV"""
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    
    # Texte en-tête
    run = header_para.add_run(f"{DIAGPV_CONFIG['nom']} | {DIAGPV_CONFIG['telephone']} | {DIAGPV_CONFIG['email']}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_secondaire'])
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def ajouter_page_garde(doc: Document, centrale_info: Dict[str, Any], data_terrain: Dict[str, Any]):
    """Ajoute page de garde rapport"""
    
    # Espacement
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Titre principal
    titre = doc.add_paragraph()
    run = titre.add_run("RAPPORT D'AUDIT PHOTOVOLTAÏQUE")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sous-titre mission
    sous_titre = doc.add_paragraph()
    run = sous_titre.add_run("Mission GIRASOLE 2025 - Audit Qualité Visuelle")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_texte'])
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Informations centrale
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    rows_data = [
        ("ID Centrale", centrale_info.get('id', 'N/A')),
        ("Nom Installation", centrale_info.get('nom', 'N/A')),
        ("Puissance", f"{centrale_info.get('puissance_kwc', 'N/A')} kWc"),
        ("Type", centrale_info.get('type', 'N/A')),
        ("Date Audit", data_terrain.get('date_audit', datetime.now().strftime('%d/%m/%Y')))
    ]
    
    for i, (label, valeur) in enumerate(rows_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(valeur)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Informations DiagPV
    footer_para = doc.add_paragraph()
    footer_para.add_run(f"{DIAGPV_CONFIG['nom']}\n").bold = True
    footer_para.add_run(f"{DIAGPV_CONFIG['adresse']}, {DIAGPV_CONFIG['code_postal']} {DIAGPV_CONFIG['ville']}\n")
    footer_para.add_run(f"☎ {DIAGPV_CONFIG['telephone']} | ✉ {DIAGPV_CONFIG['email']}\n")
    footer_para.add_run(f"{DIAGPV_CONFIG['rcs']}\n")
    footer_para.add_run(f"{DIAGPV_CONFIG['baseline']}")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for run in footer_para.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_secondaire'])
    
    # Saut de page
    doc.add_page_break()


def ajouter_synthese_executive(doc: Document, data_be: Dict[str, Any]):
    """Ajoute synthèse exécutive"""
    
    # Titre section
    titre = doc.add_heading('SYNTHÈSE EXÉCUTIVE', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    # Statut audit
    statut = data_be.get('statut_audit', 'À compléter')
    statut_para = doc.add_paragraph()
    statut_para.add_run("Statut Installation : ").bold = True
    statut_run = statut_para.add_run(statut)
    statut_run.bold = True
    
    if "conforme" in statut.lower():
        statut_run.font.color.rgb = RGBColor(46, 204, 113)  # Vert
    else:
        statut_run.font.color.rgb = RGBColor(231, 76, 60)  # Rouge
    
    # Anomalies
    doc.add_paragraph()
    anomalies_para = doc.add_paragraph()
    anomalies_para.add_run("Anomalies Détectées :\n").bold = True
    anomalies_para.add_run(f"• Critiques : {data_be.get('nb_anomalies_critiques', 0)}\n")
    anomalies_para.add_run(f"• Majeures : {data_be.get('nb_anomalies_majeures', 0)}\n")
    anomalies_para.add_run(f"• Mineures : {data_be.get('nb_anomalies_mineures', 0)}")
    
    doc.add_paragraph()


def ajouter_caracteristiques(doc: Document, centrale_info: Dict[str, Any]):
    """Ajoute caractéristiques installation"""
    
    titre = doc.add_heading('CARACTÉRISTIQUES INSTALLATION', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    # Tableau caractéristiques
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    carac_data = [
        ("Puissance Crête", f"{centrale_info.get('puissance_kwc', 'N/A')} kWc"),
        ("Modules", f"{centrale_info.get('type_modules', 'DMEGC 455Wc polycristallin')}"),
        ("Onduleurs", f"{centrale_info.get('type_onduleurs', 'DOMESOLAR / HUAWEI')}"),
        ("Type Installation", centrale_info.get('type', 'N/A')),
        ("Département", centrale_info.get('dept', 'N/A')),
        ("Mise en Service", centrale_info.get('date_mes', 'N/A'))
    ]
    
    for i, (label, valeur) in enumerate(carac_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(valeur)
    
    doc.add_paragraph()


def ajouter_methodologie(doc: Document, centrale_info: Dict[str, Any]):
    """Ajoute méthodologie audit"""
    
    titre = doc.add_heading('MÉTHODOLOGIE AUDIT', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    # Type audit
    type_install = centrale_info.get('type', 'SOL')
    
    methodo_text = """L'audit a été réalisé conformément aux normes suivantes :
• IEC 62446-1 : Inspection des installations photovoltaïques
• NF C 15-100 : Installations électriques basse tension
"""
    
    if "TOITURE" in type_install.upper():
        methodo_text += "• DTU 40.35 : Couverture par éléments porteurs photovoltaïques\n"
    
    methodo_text += """
Type d'inspection : Audit qualité visuelle
• Inspection visuelle modules et structures
• Photos standardisées nomenclature GIRASOLE (6 catégories)
• Mesures électriques basiques (Voc, Isc, isolement)
• Analyse documentaire (DOE, plans, schémas)

Méthodes NON appliquées :
• Électroluminescence (EL)
• Thermographie infrarouge drone
• Courbes I-V détaillées
"""
    
    doc.add_paragraph(methodo_text)


def ajouter_anomalies_photos(doc: Document, data_be: Dict[str, Any], photos: Dict[str, bytes]):
    """Ajoute anomalies détectées avec photos"""
    
    titre = doc.add_heading('ANOMALIES DÉTECTÉES', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    anomalies = data_be.get('anomalies_liste', [])
    
    if not anomalies:
        doc.add_paragraph("✅ Aucune anomalie majeure détectée lors de l'inspection.")
        return
    
    # Regrouper par criticité
    anomalies_critiques = [a for a in anomalies if a.get('criticite') == 'CRITIQUE']
    anomalies_majeures = [a for a in anomalies if a.get('criticite') == 'MAJEUR']
    anomalies_mineures = [a for a in anomalies if a.get('criticite') == 'MINEUR']
    
    for criticite, liste in [('CRITIQUES', anomalies_critiques), 
                              ('MAJEURES', anomalies_majeures), 
                              ('MINEURES', anomalies_mineures)]:
        if liste:
            doc.add_heading(f'Anomalies {criticite}', level=2)
            
            for i, anomalie in enumerate(liste, 1):
                # Titre anomalie
                anomalie_para = doc.add_paragraph()
                anomalie_para.add_run(f"{i}. {anomalie.get('titre', 'Anomalie')}").bold = True
                
                # Description
                doc.add_paragraph(f"Description : {anomalie.get('description', 'N/A')}")
                doc.add_paragraph(f"Localisation : {anomalie.get('localisation', 'N/A')}")
                
                # Photo si disponible
                photo_ref = anomalie.get('photo_ref')
                if photo_ref and photo_ref in photos:
                    try:
                        img_bytes = redimensionner_image(photos[photo_ref], largeur_max=350)
                        doc.add_picture(io.BytesIO(img_bytes), width=Inches(3.5))
                    except Exception as e:
                        print(f"   ⚠️  Erreur insertion photo {photo_ref}: {e}")
                
                doc.add_paragraph()


def ajouter_recommandations(doc: Document, data_be: Dict[str, Any]):
    """Ajoute recommandations priorisées"""
    
    titre = doc.add_heading('RECOMMANDATIONS', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    recommandations = data_be.get('recommandations_prioritaires', 'Aucune recommandation spécifique.')
    
    if isinstance(recommandations, list):
        for i, reco in enumerate(recommandations, 1):
            doc.add_paragraph(f"{i}. {reco}", style='List Number')
    else:
        doc.add_paragraph(recommandations)
    
    doc.add_paragraph()


def ajouter_annexes_photos(doc: Document, photos: Dict[str, bytes]):
    """Ajoute annexes avec toutes les photos par catégorie"""
    
    doc.add_page_break()
    titre = doc.add_heading('ANNEXES - PHOTOGRAPHIES', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    # Regrouper photos par catégorie GIRASOLE
    categories = {
        'GEN': 'Vues Générales',
        'ELEC': 'Installations Électriques (NF C 15-100)',
        'TOIT': 'Toiture et Étanchéité (DTU 40.35)',
        'BP': 'Bonnes Pratiques',
        'MP': 'Mauvaises Pratiques',
        'DOC': 'Documentation'
    }
    
    for cat_code, cat_nom in categories.items():
        photos_cat = {k: v for k, v in photos.items() if cat_code in k.upper()}
        
        if photos_cat:
            doc.add_heading(cat_nom, level=2)
            
            for i, (nom_photo, img_bytes) in enumerate(photos_cat.items(), 1):
                try:
                    doc.add_paragraph(f"Photo {i} : {nom_photo}")
                    img_bytes_resized = redimensionner_image(img_bytes, largeur_max=400)
                    doc.add_picture(io.BytesIO(img_bytes_resized), width=Inches(4))
                    doc.add_paragraph()
                except Exception as e:
                    print(f"   ⚠️  Erreur insertion photo annexe {nom_photo}: {e}")


def ajouter_pied_page(doc: Document):
    """Ajoute pied de page avec signature et disclaimer"""
    
    section = doc.sections[0]
    footer = section.footer
    
    # Signature
    footer_para = footer.paragraphs[0]
    footer_para.add_run(f"\n{SIGNATAIRE['nom']}\n").bold = True
    footer_para.add_run(f"{SIGNATAIRE['fonction']}\n")
    footer_para.add_run(f"{SIGNATAIRE['societe']}\n\n")
    
    # Disclaimer
    disclaimer_run = footer_para.add_run(DISCLAIMER)
    disclaimer_run.font.size = Pt(8)
    disclaimer_run.font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_secondaire'])
    
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ====================================================================================================
# 🚀 FONCTION PRINCIPALE GÉNÉRATION RAPPORT
# ====================================================================================================

def generer_rapport_pdf(centrale_id: str, json_terrain_path: Path, json_be_path: Path, 
                        photos_zip_path: Path, centrale_info: Dict[str, Any], 
                        output_dir: Path) -> Path:
    """
    Génère 1 rapport PDF branded DiagPV pour 1 centrale
    
    Args:
        centrale_id: ID centrale (ex: '3085')
        json_terrain_path: Chemin JSON checklist terrain
        json_be_path: Chemin JSON checklist BE
        photos_zip_path: Chemin ZIP photos nomenclature GIRASOLE
        centrale_info: Dict infos centrale (depuis liste_52_centrales.json)
        output_dir: Dossier sortie PDF
    
    Returns:
        Path du rapport PDF généré
    """
    
    print(f"\n📄 Génération rapport centrale {centrale_id}...")
    
    # 1. Charger données JSON
    data_terrain = charger_json(json_terrain_path)
    data_be = charger_json(json_be_path)
    
    if not data_terrain or not data_be:
        print(f"   ❌ Données manquantes pour centrale {centrale_id}")
        return None
    
    # 2. Extraire photos
    photos = extraire_photos_zip(photos_zip_path) if photos_zip_path.exists() else {}
    
    # 3. Créer document Word
    doc = Document()
    
    # Configurer marges
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # 4. Ajouter sections rapport
    ajouter_entete_diagpv(doc)
    ajouter_page_garde(doc, centrale_info, data_terrain)
    ajouter_synthese_executive(doc, data_be)
    ajouter_caracteristiques(doc, centrale_info)
    ajouter_methodologie(doc, centrale_info)
    ajouter_anomalies_photos(doc, data_be, photos)
    ajouter_recommandations(doc, data_be)
    
    if photos:
        ajouter_annexes_photos(doc, photos)
    
    ajouter_pied_page(doc)
    
    # 5. Sauvegarder Word
    nom_fichier = f"RAPPORT_{centrale_id}_{centrale_info.get('nom', 'CENTRALE').replace(' ', '_')}"
    output_docx = output_dir / f"{nom_fichier}.docx"
    doc.save(output_docx)
    
    print(f"   ✅ Rapport Word généré : {output_docx.name}")
    
    # 6. Conversion Word → PDF (optionnel, nécessite LibreOffice)
    # Pour conversion PDF, exécuter :
    # libreoffice --headless --convert-to pdf --outdir output_dir output_docx
    
    return output_docx


# ====================================================================================================
# 🎯 MAIN - BATCH PROCESSING 52 CENTRALES
# ====================================================================================================

def main():
    """Génère 52 rapports PDF"""
    
    print("="*100)
    print("📄 GÉNÉRATEUR RAPPORTS PDF INDIVIDUELS - GIRASOLE 2025")
    print("="*100)
    print()
    
    # Chemins
    base_dir = Path(__file__).parent.parent
    data_dir = base_dir / 'data'
    output_dir = base_dir / 'output_rapports'
    output_dir.mkdir(exist_ok=True)
    
    # Charger liste 52 centrales
    centrales_path = base_dir / 'templates' / 'liste_52_centrales.json'
    if not centrales_path.exists():
        print(f"❌ Fichier liste centrales introuvable : {centrales_path}")
        return
    
    centrales = charger_json(centrales_path)
    print(f"📊 {len(centrales)} centrales à traiter\n")
    
    # Statistiques
    nb_ok = 0
    nb_erreurs = 0
    
    # Générer rapport pour chaque centrale
    for i, centrale in enumerate(centrales, 1):
        centrale_id = str(centrale.get('id', ''))
        
        # Chemins fichiers
        json_terrain = data_dir / f"CHECKLIST_{centrale_id}_terrain.json"
        json_be = data_dir / f"CHECKLIST_{centrale_id}_BE.json"
        photos_zip = data_dir / f"PHOTOS_{centrale_id}.zip"
        
        # Vérifier fichiers existent
        if not json_terrain.exists() or not json_be.exists():
            print(f"[{i}/{len(centrales)}] ⚠️  SKIP {centrale_id} - Fichiers JSON manquants")
            nb_erreurs += 1
            continue
        
        try:
            # Générer rapport
            output_path = generer_rapport_pdf(
                centrale_id=centrale_id,
                json_terrain_path=json_terrain,
                json_be_path=json_be,
                photos_zip_path=photos_zip,
                centrale_info=centrale,
                output_dir=output_dir
            )
            
            if output_path:
                print(f"[{i}/{len(centrales)}] ✅ {centrale_id} - {centrale.get('nom', 'N/A')}")
                nb_ok += 1
            else:
                nb_erreurs += 1
        
        except Exception as e:
            print(f"[{i}/{len(centrales)}] ❌ ERREUR {centrale_id}: {e}")
            nb_erreurs += 1
    
    # Statistiques finales
    print()
    print("="*100)
    print("📊 STATISTIQUES GÉNÉRATION")
    print("="*100)
    print(f"✅ Rapports générés : {nb_ok}/{len(centrales)}")
    print(f"❌ Erreurs : {nb_erreurs}")
    print(f"📁 Dossier sortie : {output_dir}")
    print("="*100)
    print()
    print("💡 NOTE : Pour conversion PDF, exécuter :")
    print(f"   libreoffice --headless --convert-to pdf --outdir {output_dir} {output_dir}/*.docx")
    print()


if __name__ == "__main__":
    main()
