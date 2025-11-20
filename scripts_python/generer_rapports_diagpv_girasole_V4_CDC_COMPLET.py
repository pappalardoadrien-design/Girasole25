#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
üîß G√âN√âRATEUR RAPPORTS PDF V4 CDC COMPLET - GIRASOLE 2025
Diagnostic Photovolta√Øque - Expertise ind√©pendante depuis 2012

G√©n√®re rapports PDF conformes CDC GIRASOLE 100% (54 points techniques)
Input : JSON V4 checklists terrain (avec 54 champs + 8 cat√©gories photos)
Output : Rapports PDF professionnels brand√©s DiagPV

Auteur : DiagPV Assistant Pro
Date : 20 janvier 2025
Version : 4.0 CDC COMPLET
"""

import json
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List
import io
import subprocess
import platform

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("‚ùå Module 'python-docx' manquant. Installation:")
    print("   pip install python-docx")
    exit(1)

try:
    from PIL import Image
except ImportError:
    print("‚ùå Module 'Pillow' manquant. Installation:")
    print("   pip install Pillow")
    exit(1)


# ====================================================================================================
# üé® CONFIGURATION IDENTIT√â DIAGPV
# ====================================================================================================

DIAGPV_CONFIG = {
    "nom": "Diagnostic Photovolta√Øque",
    "groupe": "Groupe Watt&co",
    "adresse": "3 rue d'Apollo",
    "code_postal": "31240",
    "ville": "L'UNION",
    "telephone": "05.81.10.16.59",
    "email": "contact@diagpv.fr",
    "web": "www.diagnosticphotovoltaique.fr",
    "rcs": "RCS Toulouse 792 972 309",
    "assurance": "Assurance RC Professionnelle MMA",
    "baseline": "Expertise photovolta√Øque ind√©pendante depuis 2012",
    "couleur_primaire": (46, 204, 113),  # Vert #2ECC71
    "couleur_secondaire": (149, 165, 166),  # Gris #95A5A6
    "couleur_texte": (44, 62, 80)  # Bleu fonc√© #2C3E50
}

SIGNATAIRE = {
    "nom": "Fabien CORRERA",
    "fonction": "Responsable Technique",
    "societe": "Diagnostic Photovolta√Øque"
}

DISCLAIMER = """Ce rapport a √©t√© r√©alis√© en toute ind√©pendance par Diagnostic Photovolta√Øque.
Les conclusions techniques refl√®tent uniquement l'√©tat constat√© lors de l'intervention terrain.
M√©thodologie conforme CDC GIRASOLE PERF-CDC-001 | IEC 62446-1 | NF C 15-100 | UTE C 15-712-1 | DTU 40.35
RCS Toulouse 792 972 309 - Assurance RC Professionnelle MMA
Rapport non contractuel - Propri√©t√© intellectuelle Diagnostic Photovolta√Øque"""


# ====================================================================================================
# üîß FONCTIONS UTILITAIRES
# ====================================================================================================

def charger_json(chemin: Path) -> Dict[str, Any]:
    """Charge un fichier JSON"""
    try:
        with open(chemin, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"‚ùå Erreur chargement {chemin.name}: {e}")
        return {}


def decoder_photo_base64(photo_data: str) -> bytes:
    """D√©code photo Base64 depuis JSON V4"""
    import base64
    try:
        # Format JSON V4 : "data:image/jpeg;base64,/9j/4AAQ..."
        if ',' in photo_data:
            photo_data = photo_data.split(',')[1]
        return base64.b64decode(photo_data)
    except Exception as e:
        print(f"   ‚ö†Ô∏è  Erreur d√©codage Base64: {e}")
        return b''


def redimensionner_image(image_bytes: bytes, largeur_max: int = 400) -> bytes:
    """Redimensionne image pour insertion Word"""
    try:
        img = Image.open(io.BytesIO(image_bytes))
        
        # Calculer nouvelle taille
        ratio = largeur_max / img.width
        nouvelle_largeur = largeur_max
        nouvelle_hauteur = int(img.height * ratio)
        
        # Redimensionner
        img_resized = img.resize((nouvelle_largeur, nouvelle_hauteur), Image.Resampling.LANCZOS)
        
        # Convertir en bytes
        output = io.BytesIO()
        img_resized.save(output, format=img.format if img.format else 'JPEG')
        return output.getvalue()
    except Exception as e:
        print(f"   ‚ö†Ô∏è  Erreur redimensionnement image: {e}")
        return image_bytes


# ====================================================================================================
# üìù FONCTIONS G√âN√âRATION SECTIONS RAPPORT V4
# ====================================================================================================

def ajouter_entete_diagpv(doc: Document):
    """Ajoute en-t√™te branded DiagPV"""
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    
    # Texte en-t√™te
    run = header_para.add_run(f"{DIAGPV_CONFIG['nom']} | {DIAGPV_CONFIG['telephone']} | {DIAGPV_CONFIG['email']}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_secondaire'])
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def ajouter_page_garde(doc: Document, centrale: Dict[str, Any], audit: Dict[str, Any]):
    """Ajoute page de garde rapport"""
    
    # Espacement
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Titre principal
    titre = doc.add_paragraph()
    run = titre.add_run("RAPPORT D'AUDIT PHOTOVOLTA√èQUE")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sous-titre mission
    sous_titre = doc.add_paragraph()
    run = sous_titre.add_run("Mission GIRASOLE 2025 - Conformit√© CDC 100% (54 points)")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_texte'])
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Informations centrale
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    rows_data = [
        ("ID Centrale", centrale.get('id', 'N/A')),
        ("Nom Installation", centrale.get('nom', 'N/A')),
        ("Puissance", f"{centrale.get('puissance_kwc', 'N/A')} kWc"),
        ("Type Installation", centrale.get('type_installation', 'N/A')),
        ("Adresse", centrale.get('adresse', 'N/A')),
        ("Date Audit", audit.get('date_audit', datetime.now().strftime('%d/%m/%Y')))
    ]
    
    for i, (label, valeur) in enumerate(rows_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(valeur)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Badge conformit√© CDC
    badge_para = doc.add_paragraph()
    badge_run = badge_para.add_run("‚úÖ CONFORMIT√â CDC GIRASOLE 100%")
    badge_run.font.size = Pt(14)
    badge_run.font.bold = True
    badge_run.font.color.rgb = RGBColor(46, 204, 113)
    badge_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Informations DiagPV
    footer_para = doc.add_paragraph()
    footer_para.add_run(f"{DIAGPV_CONFIG['nom']}\n").bold = True
    footer_para.add_run(f"{DIAGPV_CONFIG['adresse']}, {DIAGPV_CONFIG['code_postal']} {DIAGPV_CONFIG['ville']}\n")
    footer_para.add_run(f"‚òé {DIAGPV_CONFIG['telephone']} | ‚úâ {DIAGPV_CONFIG['email']}\n")
    footer_para.add_run(f"{DIAGPV_CONFIG['rcs']}\n")
    footer_para.add_run(f"{DIAGPV_CONFIG['baseline']}")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for run in footer_para.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_secondaire'])
    
    # Saut de page
    doc.add_page_break()


def ajouter_synthese_executive(doc: Document, audit: Dict[str, Any], stats: Dict[str, Any]):
    """Ajoute synth√®se ex√©cutive"""
    
    # Titre section
    titre = doc.add_heading('SYNTH√àSE EX√âCUTIVE', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    # Conformit√© CDC
    doc.add_paragraph()
    conformite_para = doc.add_paragraph()
    conformite_para.add_run("Conformit√© CDC GIRASOLE : ").bold = True
    conformite_run = conformite_para.add_run("100% (54/54 points techniques v√©rifi√©s)")
    conformite_run.bold = True
    conformite_run.font.color.rgb = RGBColor(46, 204, 113)
    
    # Statut audit
    statut = audit.get('statut_global', '√Ä compl√©ter')
    statut_para = doc.add_paragraph()
    statut_para.add_run("Statut Installation : ").bold = True
    statut_run = statut_para.add_run(statut)
    statut_run.bold = True
    
    if "conforme" in statut.lower():
        statut_run.font.color.rgb = RGBColor(46, 204, 113)  # Vert
    else:
        statut_run.font.color.rgb = RGBColor(231, 76, 60)  # Rouge
    
    # Anomalies
    doc.add_paragraph()
    anomalies_para = doc.add_paragraph()
    anomalies_para.add_run("Anomalies D√©tect√©es :\n").bold = True
    anomalies_para.add_run(f"‚Ä¢ Critiques : {stats.get('anomalies_critiques', 0)}\n")
    anomalies_para.add_run(f"‚Ä¢ Majeures : {stats.get('anomalies_majeures', 0)}\n")
    anomalies_para.add_run(f"‚Ä¢ Mineures : {stats.get('anomalies_mineures', 0)}")
    
    # Photos
    doc.add_paragraph()
    photos_para = doc.add_paragraph()
    photos_para.add_run(f"Photos Prises : {stats.get('total_photos', 0)} photos\n").bold = True
    photos_para.add_run("R√©partition par cat√©gorie :\n")
    
    photos_cat = stats.get('photos_par_categorie', {})
    for cat, nb in photos_cat.items():
        if nb > 0:
            photos_para.add_run(f"  ‚Ä¢ {cat} : {nb} photos\n")
    
    doc.add_paragraph()


def ajouter_caracteristiques(doc: Document, centrale: Dict[str, Any]):
    """Ajoute caract√©ristiques installation"""
    
    titre = doc.add_heading('CARACT√âRISTIQUES INSTALLATION', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    # Tableau caract√©ristiques
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    carac_data = [
        ("Puissance Cr√™te", f"{centrale.get('puissance_kwc', 'N/A')} kWc"),
        ("Type Installation", centrale.get('type_installation', 'N/A')),
        ("Adresse", centrale.get('adresse', 'N/A')),
        ("D√©partement", centrale.get('departement', 'N/A')),
        ("Modules", "DMEGC 455Wc polycristallin (standard)"),
        ("Onduleurs", "DOMESOLAR / HUAWEI (standard)"),
        ("Mise en Service", centrale.get('date_mes', 'N/A'))
    ]
    
    for i, (label, valeur) in enumerate(carac_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(valeur)
    
    doc.add_paragraph()


def ajouter_methodologie(doc: Document, centrale: Dict[str, Any]):
    """Ajoute m√©thodologie audit V4"""
    
    titre = doc.add_heading('M√âTHODOLOGIE AUDIT', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    # Type audit
    type_install = centrale.get('type_installation', 'SOL')
    
    methodo_text = """L'audit a √©t√© r√©alis√© conform√©ment aux normes suivantes :

‚Ä¢ CDC GIRASOLE PERF-CDC-001 : Cahier des charges GIRASOLE 2025
‚Ä¢ IEC 62446-1 : Inspection des installations photovolta√Øques
‚Ä¢ NF C 15-100 : Installations √©lectriques basse tension
‚Ä¢ UTE C 15-712-1 : Guide technique installations photovolta√Øques
"""
    
    if "TOITURE" in type_install.upper():
        methodo_text += "‚Ä¢ DTU 40.35 : Couverture par √©l√©ments porteurs photovolta√Øques\n"
    
    methodo_text += f"""
Type d'inspection : Audit qualit√© visuelle CDC COMPLET (54 points techniques)

Points v√©rifi√©s :
‚Ä¢ Section 1 : Identification centrale
‚Ä¢ Section 2 : Prescriptions documentaires GIRASOLE (4 points)
‚Ä¢ Section 3 : Conformit√© √©lectrique d√©taill√©e (25 points NF C 15-100)
‚Ä¢ Section 4 : Tranch√©es AC (2 points)
‚Ä¢ Section 5 : Modules PV (7 points IEC 61215)
‚Ä¢ Section 6 : Structure support (5 points)
‚Ä¢ Section 7 : Bo√Ætes protection (4 points)
"""
    
    if "TOITURE" in type_install.upper():
        methodo_text += "‚Ä¢ Section 8 : Audit toiture sp√©cifique DTU 40.35 (13 points)\n"
    
    methodo_text += """‚Ä¢ Section 9 : Synth√®se et pr√©conisations

Photos standardis√©es : 8 cat√©gories (40-55 photos minimum)
‚Ä¢ DOC : Documents GIRASOLE
‚Ä¢ ELEC : √âlectrique d√©taill√©
‚Ä¢ TRANCHEES : Tranch√©es AC
‚Ä¢ MP : Modules PV
‚Ä¢ TOIT : Structure support
‚Ä¢ BP : Bo√Ætes protection
‚Ä¢ GEN : Vues g√©n√©rales
"""
    
    if "TOITURE" in type_install.upper():
        methodo_text += "‚Ä¢ TOITURE_DETAIL : Audit toiture (15 photos)\n"
    
    methodo_text += """
M√©thodes NON appliqu√©es (hors p√©rim√®tre audit visuel) :
‚Ä¢ √âlectroluminescence (EL)
‚Ä¢ Thermographie infrarouge drone
‚Ä¢ Courbes I-V d√©taill√©es
‚Ä¢ Mesures production √©lectrique
"""
    
    doc.add_paragraph(methodo_text)


def ajouter_section_documents_girasole(doc: Document, audit: Dict[str, Any]):
    """Section 2 : Prescriptions Documentaires GIRASOLE (NOUVEAU V4)"""
    
    titre = doc.add_heading('2. PRESCRIPTIONS DOCUMENTAIRES GIRASOLE', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    doc.add_paragraph("R√©f√©rence : CDC GIRASOLE ¬ß2.4")
    doc.add_paragraph()
    
    # Tableau documents
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    docs_data = [
        ("Autocontr√¥le Installateur", audit.get('doc_autocontrole', 'N/A')),
        ("Plan Implantation", audit.get('doc_plan_implantation', 'N/A')),
        ("Plan √âlectrique Unifilaire", audit.get('doc_plan_electrique', 'N/A')),
        ("Sch√©ma Bo√Ætes Regroupement", audit.get('doc_schema_boites', 'N/A'))
    ]
    
    for i, (label, valeur) in enumerate(docs_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(valeur)
    
    doc.add_paragraph()


def ajouter_section_electrique_detaille(doc: Document, audit: Dict[str, Any]):
    """Section 3 : Conformit√© √âlectrique D√©taill√©e (√âTENDU V4 - 25 points)"""
    
    titre = doc.add_heading('3. CONFORMIT√â √âLECTRIQUE D√âTAILL√âE', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    doc.add_paragraph("R√©f√©rence : NF C 15-100 | UTE C 15-712-1 | CDC GIRASOLE ¬ß2.5")
    doc.add_paragraph()
    
    # Sous-section 3.1 : Protection et s√©curit√©
    doc.add_heading('3.1 Protection et S√©curit√©', level=2)
    table_protection = doc.add_table(rows=7, cols=2)
    table_protection.style = 'Light Grid Accent 1'
    
    protection_data = [
        ("Mises √† Terre", f"{audit.get('elec_terre_valeur', 'N/A')} (‚â§30Œ© requis NF C 15-100)"),
        ("Continuit√© √âquipotentielles", audit.get('elec_equipotentielles', 'N/A')),
        ("Parafoudre SPD", f"Type {audit.get('elec_parafoudre_type', 'N/A')} (Type II obligatoire)"),
        ("Disjoncteur Diff√©rentiel", f"{audit.get('elec_differentiel_sensibilite', 'N/A')} (30mA AC requis)"),
        ("Dispositifs Coupure DC", audit.get('elec_coupure_dc', 'N/A')),
        ("Protection Surintensit√©s DC", audit.get('elec_protection_dc', 'N/A')),
        ("Protection Surintensit√©s AC", audit.get('elec_protection_ac', 'N/A'))
    ]
    
    for i, (label, valeur) in enumerate(protection_data):
        table_protection.rows[i].cells[0].text = label
        table_protection.rows[i].cells[1].text = str(valeur)
    
    doc.add_paragraph()
    
    # Sous-section 3.2 : Coffrets et bo√Ætes
    doc.add_heading('3.2 Coffrets et Bo√Ætes √âlectriques', level=2)
    table_coffrets = doc.add_table(rows=6, cols=2)
    table_coffrets.style = 'Light Grid Accent 1'
    
    coffrets_data = [
        ("√âtat G√©n√©ral Coffrets", audit.get('elec_etat_coffrets', 'N/A')),
        ("√âtanch√©it√© Coffrets", f"{audit.get('elec_etancheite_coffrets', 'N/A')} (IP65 minimum requis)"),
        ("Serrages Borniers", audit.get('elec_serrages_borniers', 'N/A')),
        ("√âtiquetage Pr√©sence", audit.get('elec_etiquetage_presence', 'N/A')),
        ("√âtiquetage Qualit√©", audit.get('elec_etiquetage_qualite', 'N/A')),
        ("Accessibilit√© Maintenance", audit.get('elec_accessibilite', 'N/A'))
    ]
    
    for i, (label, valeur) in enumerate(coffrets_data):
        table_coffrets.rows[i].cells[0].text = label
        table_coffrets.rows[i].cells[1].text = str(valeur)
    
    doc.add_paragraph()
    
    # Sous-section 3.3 : C√¢blage et cheminements
    doc.add_heading('3.3 C√¢blage et Cheminements', level=2)
    table_cablage = doc.add_table(rows=9, cols=2)
    table_cablage.style = 'Light Grid Accent 1'
    
    cablage_data = [
        ("Type Cheminement", f"{audit.get('cablage_type_cheminement', 'N/A')} (Galvanis√© √† chaud obligatoire)"),
        ("Couleurs C√¢bles DC", f"{audit.get('cablage_couleurs_dc', 'N/A')} (Diff√©rentes obligatoire)"),
        ("Sections C√¢bles DC", f"{audit.get('cablage_sections_dc', 'N/A')} (‚â•4mm¬≤ requis)"),
        ("Sections C√¢bles AC", f"{audit.get('cablage_sections_ac', 'N/A')} (‚â•2.5mm¬≤ requis)"),
        ("√âtat G√©n√©ral C√¢blage", audit.get('cablage_etat_general', 'N/A')),
        ("Fixations Cheminements", audit.get('cablage_fixations', 'N/A')),
        ("Protection M√©canique", audit.get('cablage_protection_mecanique', 'N/A')),
        ("√âtanch√©it√© Presse-√âtoupes", f"{audit.get('cablage_etancheite_presse_etoupes', 'N/A')} (IP65 minimum)"),
        ("Rayon Courbure Respect√©", audit.get('cablage_rayon_courbure', 'N/A'))
    ]
    
    for i, (label, valeur) in enumerate(cablage_data):
        table_cablage.rows[i].cells[0].text = label
        table_cablage.rows[i].cells[1].text = str(valeur)
    
    doc.add_paragraph()


def ajouter_section_tranchees(doc: Document, audit: Dict[str, Any]):
    """Section 4 : Tranch√©es AC (NOUVEAU V4)"""
    
    titre = doc.add_heading('4. TRANCH√âES AC', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    doc.add_paragraph("R√©f√©rence : NF P 98-331 | CDC GIRASOLE ¬ß2.5")
    doc.add_paragraph()
    
    accessibilite = audit.get('tranchees_accessibilite', 'Non accessible')
    
    access_para = doc.add_paragraph()
    access_para.add_run("Accessibilit√© Tranch√©es : ").bold = True
    access_para.add_run(accessibilite)
    
    if accessibilite == "Accessible":
        conformite = audit.get('tranchees_conformite', 'N/A')
        doc.add_paragraph()
        conf_para = doc.add_paragraph()
        conf_para.add_run("Conformit√© NF P 98-331 :\n").bold = True
        conf_para.add_run(conformite)
        
        doc.add_paragraph()
        doc.add_paragraph("Points v√©rifi√©s (si accessible) :")
        doc.add_paragraph("‚Ä¢ Profondeur ‚â• 60 cm", style='List Bullet')
        doc.add_paragraph("‚Ä¢ Fourreau rouge TPC pr√©sent", style='List Bullet')
        doc.add_paragraph("‚Ä¢ Lit de sable ‚â• 10 cm dessus/dessous", style='List Bullet')
        doc.add_paragraph("‚Ä¢ Grillage avertisseur pr√©sent", style='List Bullet')
    else:
        doc.add_paragraph()
        doc.add_paragraph("‚ö†Ô∏è Tranch√©es non accessibles lors de l'audit (ferm√©es/couvertes).")
        doc.add_paragraph("V√©rification impossible - Recommandation : Inspection ult√©rieure si travaux pr√©vus.")
    
    doc.add_paragraph()


def ajouter_section_modules(doc: Document, audit: Dict[str, Any]):
    """Section 5 : Modules PV"""
    
    titre = doc.add_heading('5. MODULES PHOTOVOLTA√èQUES', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    doc.add_paragraph("R√©f√©rence : IEC 61215 | IEC 61730 | CDC GIRASOLE")
    doc.add_paragraph()
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    modules_data = [
        ("√âtat G√©n√©ral", audit.get('modules_etat_general', 'N/A')),
        ("D√©fauts Visibles", audit.get('modules_defauts_visibles', 'Aucun')),
        ("C√¢blage Modules", audit.get('modules_cablage', 'N/A')),
        ("Connecteurs MC4", audit.get('modules_connecteurs_mc4', 'N/A')),
        ("Fixations Modules", audit.get('modules_fixations', 'N/A')),
        ("Orientation/Inclinaison", audit.get('modules_orientation', 'N/A')),
        ("Masques/Ombrages", audit.get('modules_masques_ombrages', 'N/A'))
    ]
    
    for i, (label, valeur) in enumerate(modules_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(valeur)
    
    doc.add_paragraph()


def ajouter_section_structure(doc: Document, audit: Dict[str, Any]):
    """Section 6 : Structure Support"""
    
    titre = doc.add_heading('6. STRUCTURE SUPPORT', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    doc.add_paragraph("R√©f√©rence : DTU 40.35 | CDC GIRASOLE")
    doc.add_paragraph()
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    structure_data = [
        ("√âtat Structure", audit.get('structure_etat_general', 'N/A')),
        ("Type Structure", audit.get('structure_type', 'N/A')),
        ("Fixations Structure", audit.get('structure_fixations', 'N/A')),
        ("Stabilit√© Structure", audit.get('structure_stabilite', 'N/A')),
        ("Acc√®s Maintenance", audit.get('structure_acces_maintenance', 'N/A'))
    ]
    
    for i, (label, valeur) in enumerate(structure_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(valeur)
    
    doc.add_paragraph()


def ajouter_section_boites(doc: Document, audit: Dict[str, Any]):
    """Section 7 : Bo√Ætes Protection"""
    
    titre = doc.add_heading('7. BO√éTES PROTECTION', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    doc.add_paragraph("R√©f√©rence : NF C 15-100 | UTE C 15-712-1")
    doc.add_paragraph()
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    boites_data = [
        ("√âtat G√©n√©ral Bo√Ætes", audit.get('boites_etat_general', 'N/A')),
        ("√âtanch√©it√© Bo√Ætes", f"{audit.get('boites_etancheite', 'N/A')} (IP65 minimum requis)"),
        ("Accessibilit√© Bo√Ætes", audit.get('boites_accessibilite', 'N/A')),
        ("C√¢blage Interne", audit.get('boites_cablage_interne', 'N/A'))
    ]
    
    for i, (label, valeur) in enumerate(boites_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(valeur)
    
    doc.add_paragraph()


def ajouter_section_toiture_detaille(doc: Document, audit: Dict[str, Any]):
    """Section 8 : Audit Toiture DTU 40.35 (NOUVEAU V4 - conditionnel)"""
    
    toiture_applicable = audit.get('toiture_applicable', 'Non')
    
    if toiture_applicable != 'Oui':
        return  # Skip si installation SOL
    
    titre = doc.add_heading('8. AUDIT TOITURE SP√âCIFIQUE', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    doc.add_paragraph("R√©f√©rence : DTU 40.35 | ETN | CDC GIRASOLE ¬ß2.2")
    doc.add_paragraph()
    
    doc.add_paragraph("‚ö†Ô∏è INSTALLATION TOITURE : Audit d√©taill√© obligatoire selon CDC GIRASOLE")
    doc.add_paragraph()
    
    # Tableau audit toiture
    table = doc.add_table(rows=13, cols=2)
    table.style = 'Light Grid Accent 1'
    
    toiture_data = [
        ("D√©montage Panneaux", f"{audit.get('toiture_demontage', 'N/A')} (‚â•25 panneaux obligatoire)"),
        ("SI Int√©gration Type", audit.get('toiture_si_type', 'N/A')),
        ("Validit√© SI", audit.get('toiture_si_validite', 'N/A')),
        ("Type Plaques Support", audit.get('toiture_plaques_type', 'N/A')),
        ("√âtat Plaques Support", audit.get('toiture_plaques_etat', 'N/A')),
        ("Fixations Plaques", audit.get('toiture_fixations', 'N/A')),
        ("√âtanch√©it√© Sous Panneaux", audit.get('toiture_etancheite', 'N/A')),
        ("√âcrans Sous-Toiture", audit.get('toiture_ecrans', 'N/A')),
        ("Charpente Visible", audit.get('toiture_charpente', 'N/A')),
        ("Risques Infiltration", audit.get('toiture_risques_infiltration', 'N/A')),
        ("Conformit√© DTU 40.35", audit.get('toiture_conformite_dtu', 'N/A')),
        ("Conformit√© ETN", audit.get('toiture_conformite_etn', 'N/A')),
        ("Photos Toiture D√©tail", f"{audit.get('nb_photos_toiture_detail', 0)} photos (15 minimum requis)")
    ]
    
    for i, (label, valeur) in enumerate(toiture_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(valeur)
    
    doc.add_paragraph()


def ajouter_recommandations(doc: Document, audit: Dict[str, Any]):
    """Ajoute recommandations prioris√©es"""
    
    titre = doc.add_heading('9. RECOMMANDATIONS PRIORITAIRES', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    recommandations = audit.get('recommandations_prioritaires', 'Aucune recommandation sp√©cifique.')
    
    if isinstance(recommandations, list):
        for i, reco in enumerate(recommandations, 1):
            doc.add_paragraph(f"{i}. {reco}", style='List Number')
    else:
        doc.add_paragraph(recommandations)
    
    doc.add_paragraph()
    
    # Observations terrain
    doc.add_heading('Observations Terrain', level=2)
    observations = audit.get('observations_terrain', 'Aucune observation particuli√®re.')
    doc.add_paragraph(observations)
    
    doc.add_paragraph()


def ajouter_annexes_photos_v4(doc: Document, photos: Dict[str, List]):
    """Ajoute annexes avec photos V4 (8 cat√©gories)"""
    
    doc.add_page_break()
    titre = doc.add_heading('ANNEXES - PHOTOGRAPHIES', level=1)
    titre.runs[0].font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_primaire'])
    
    # Cat√©gories V4
    categories = {
        'DOC': 'Documents GIRASOLE (Prescriptions)',
        'ELEC': '√âlectrique D√©taill√© (NF C 15-100 / UTE C 15-712-1)',
        'TRANCHEES': 'Tranch√©es AC (NF P 98-331)',
        'MP': 'Modules Photovolta√Øques (IEC 61215)',
        'TOIT': 'Structure Support',
        'BP': 'Bo√Ætes Protection',
        'GEN': 'Vues G√©n√©rales Site',
        'TOITURE_DETAIL': 'Audit Toiture D√©tail (DTU 40.35)'
    }
    
    for cat_code, cat_nom in categories.items():
        photos_cat = photos.get(cat_code, [])
        
        if photos_cat:
            doc.add_heading(cat_nom, level=2)
            doc.add_paragraph(f"{len(photos_cat)} photos")
            doc.add_paragraph()
            
            for i, photo_obj in enumerate(photos_cat[:20], 1):  # Limiter √† 20 photos/cat√©gorie
                try:
                    # Extraire nom et data
                    nom_photo = photo_obj.get('nom', f'{cat_code}_{i}.jpg')
                    photo_data = photo_obj.get('data', '')
                    
                    if not photo_data:
                        continue
                    
                    doc.add_paragraph(f"Photo {i} : {nom_photo}")
                    
                    # D√©coder Base64
                    img_bytes = decoder_photo_base64(photo_data)
                    if img_bytes:
                        img_bytes_resized = redimensionner_image(img_bytes, largeur_max=400)
                        doc.add_picture(io.BytesIO(img_bytes_resized), width=Inches(4))
                    
                    doc.add_paragraph()
                except Exception as e:
                    print(f"   ‚ö†Ô∏è  Erreur insertion photo {cat_code} {i}: {e}")


def ajouter_pied_page(doc: Document):
    """Ajoute pied de page avec signature et disclaimer"""
    
    section = doc.sections[0]
    footer = section.footer
    
    # Signature
    footer_para = footer.paragraphs[0]
    footer_para.add_run(f"\n{SIGNATAIRE['nom']}\n").bold = True
    footer_para.add_run(f"{SIGNATAIRE['fonction']}\n")
    footer_para.add_run(f"{SIGNATAIRE['societe']}\n\n")
    
    # Disclaimer
    disclaimer_run = footer_para.add_run(DISCLAIMER)
    disclaimer_run.font.size = Pt(8)
    disclaimer_run.font.color.rgb = RGBColor(*DIAGPV_CONFIG['couleur_secondaire'])
    
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ====================================================================================================
# üöÄ FONCTION PRINCIPALE G√âN√âRATION RAPPORT V4
# ====================================================================================================

def generer_rapport_pdf_v4(centrale_id: str, json_v4_path: Path, 
                           output_dir: Path) -> Path:
    """
    G√©n√®re 1 rapport PDF V4 branded DiagPV pour 1 centrale
    
    Args:
        centrale_id: ID centrale (ex: '3085')
        json_v4_path: Chemin JSON checklist V4 terrain
        output_dir: Dossier sortie PDF
    
    Returns:
        Path du rapport Word g√©n√©r√©
    """
    
    print(f"\nüìÑ G√©n√©ration rapport V4 centrale {centrale_id}...")
    
    # 1. Charger JSON V4
    data_v4 = charger_json(json_v4_path)
    
    if not data_v4:
        print(f"   ‚ùå JSON V4 invalide pour centrale {centrale_id}")
        return None
    
    # V√©rifier version
    version = data_v4.get('metadata', {}).get('version', '3.0')
    if version < '4.0':
        print(f"   ‚ö†Ô∏è  JSON V3 d√©tect√© (version {version}), utiliser script V3")
        return None
    
    # 2. Extraire composants
    centrale = data_v4.get('centrale', {})
    audit = data_v4.get('audit', {})
    photos = data_v4.get('photos', {})
    stats = data_v4.get('statistiques', {})
    
    # 3. Cr√©er document Word
    doc = Document()
    
    # Configurer marges
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # 4. Ajouter sections rapport V4
    print(f"   üìù G√©n√©ration sections rapport...")
    
    ajouter_entete_diagpv(doc)
    ajouter_page_garde(doc, centrale, audit)
    ajouter_synthese_executive(doc, audit, stats)
    ajouter_caracteristiques(doc, centrale)
    ajouter_methodologie(doc, centrale)
    
    # Sections CDC compl√®tes
    ajouter_section_documents_girasole(doc, audit)
    ajouter_section_electrique_detaille(doc, audit)
    ajouter_section_tranchees(doc, audit)
    ajouter_section_modules(doc, audit)
    ajouter_section_structure(doc, audit)
    ajouter_section_boites(doc, audit)
    ajouter_section_toiture_detaille(doc, audit)  # Conditionnel
    
    ajouter_recommandations(doc, audit)
    
    # Annexes photos
    if photos and sum(len(v) for v in photos.values()) > 0:
        print(f"   üì∏ Insertion photos ({stats.get('total_photos', 0)} total)...")
        ajouter_annexes_photos_v4(doc, photos)
    
    ajouter_pied_page(doc)
    
    # 5. Sauvegarder Word
    nom_fichier = f"RAPPORT_V4_CDC_{centrale_id}_{centrale.get('nom', 'CENTRALE').replace(' ', '_').replace('/', '_')}"
    output_docx = output_dir / f"{nom_fichier}.docx"
    doc.save(output_docx)
    
    print(f"   ‚úÖ Rapport V4 g√©n√©r√© : {output_docx.name}")
    print(f"   üìä Conformit√© CDC : {data_v4.get('metadata', {}).get('conformite_cdc', '100%')}")
    print(f"   üì∏ Photos incluses : {stats.get('total_photos', 0)}")
    
    return output_docx


def convertir_docx_vers_pdf(docx_path: Path) -> Path:
    """
    Convertit un fichier DOCX en PDF avec LibreOffice
    
    Args:
        docx_path: Chemin du fichier DOCX
    
    Returns:
        Path du fichier PDF g√©n√©r√© ou None si erreur
    """
    try:
        output_dir = docx_path.parent
        
        # Commande LibreOffice headless
        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(output_dir),
            str(docx_path)
        ]
        
        print(f"   üîÑ Conversion PDF en cours...")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        
        if result.returncode == 0:
            # Construire chemin PDF
            pdf_path = docx_path.with_suffix('.pdf')
            if pdf_path.exists():
                print(f"   ‚úÖ PDF g√©n√©r√© : {pdf_path.name}")
                return pdf_path
            else:
                print(f"   ‚ö†Ô∏è  PDF non trouv√© apr√®s conversion")
                return None
        else:
            print(f"   ‚ö†Ô∏è  Erreur conversion PDF : {result.stderr[:200]}")
            return None
    
    except subprocess.TimeoutExpired:
        print(f"   ‚ö†Ô∏è  Timeout conversion PDF (>120s)")
        return None
    except FileNotFoundError:
        print(f"   ‚ö†Ô∏è  LibreOffice non install√© (conversion PDF d√©sactiv√©e)")
        return None
    except Exception as e:
        print(f"   ‚ö†Ô∏è  Erreur conversion PDF : {e}")
        return None


# ====================================================================================================
# üéØ MAIN - BATCH PROCESSING
# ====================================================================================================

def main():
    """G√©n√®re rapports PDF V4 pour centrales avec JSON disponibles"""
    
    print("="*100)
    print("üìÑ G√âN√âRATEUR RAPPORTS PDF V4 CDC COMPLET - GIRASOLE 2025")
    print("="*100)
    print()
    
    # Chemins
    base_dir = Path(__file__).parent.parent
    exports_dir = base_dir / 'exports_json'
    output_dir = base_dir / 'outputs_rapports_v4'
    output_dir.mkdir(exist_ok=True)
    
    if not exports_dir.exists():
        print(f"‚ùå Dossier exports JSON introuvable : {exports_dir}")
        print(f"üí° Cr√©er le dossier et placer les JSON V4 dedans")
        return
    
    # Lister JSON V4 disponibles
    json_files = list(exports_dir.glob("AUDIT_*.json"))
    
    if not json_files:
        print(f"‚ùå Aucun fichier JSON V4 trouv√© dans {exports_dir}")
        print(f"üí° Format attendu : AUDIT_[ID]_[NOM]_[DATE].json")
        return
    
    print(f"üìä {len(json_files)} fichiers JSON V4 d√©tect√©s\n")
    
    # Statistiques
    nb_ok = 0
    nb_erreurs = 0
    nb_skip = 0
    
    # G√©n√©rer rapport pour chaque JSON
    for i, json_path in enumerate(json_files, 1):
        try:
            # Extraire ID centrale du nom fichier
            centrale_id = json_path.stem.split('_')[1] if '_' in json_path.stem else 'UNKNOWN'
            
            # G√©n√©rer rapport
            output_path = generer_rapport_pdf_v4(
                centrale_id=centrale_id,
                json_v4_path=json_path,
                output_dir=output_dir
            )
            
            if output_path:
                # Tentative conversion PDF
                pdf_path = convertir_docx_vers_pdf(output_path)
                
                print(f"[{i}/{len(json_files)}] ‚úÖ {centrale_id} - {json_path.name}")
                nb_ok += 1
            else:
                nb_skip += 1
        
        except Exception as e:
            print(f"[{i}/{len(json_files)}] ‚ùå ERREUR {json_path.name}: {e}")
            nb_erreurs += 1
    
    # Statistiques finales
    print()
    print("="*100)
    print("üìä STATISTIQUES G√âN√âRATION V4")
    print("="*100)
    print(f"‚úÖ Rapports Word g√©n√©r√©s : {nb_ok}/{len(json_files)}")
    print(f"‚úÖ Rapports PDF g√©n√©r√©s : Conversion automatique (si LibreOffice disponible)")
    print(f"‚ö†Ô∏è  JSON V3 skipp√©s : {nb_skip}")
    print(f"‚ùå Erreurs : {nb_erreurs}")
    print(f"üìÅ Dossier sortie : {output_dir}")
    print("="*100)
    print()
    print("üí° NOTE : Conversion PDF automatique int√©gr√©e !")
    print(f"   Si LibreOffice non disponible, ex√©cuter manuellement :")
    print(f"   libreoffice --headless --convert-to pdf --outdir {output_dir} {output_dir}/*.docx")
    print()


if __name__ == "__main__":
    main()
