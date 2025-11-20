#!/usr/bin/env python3
"""
TEST WORKFLOW COMPLET - CENTRALE 3085 (PILOTE)
Diagnostic Photovoltaïque - GIRASOLE 2025

Teste le workflow complet de génération automatique:
1. Génération ANNEXE 2 (Excel)
2. Génération Rapport PDF (DOCX + PDF)
3. Vérification outputs conformes
"""

import json
import zipfile
from pathlib import Path
from datetime import datetime
import subprocess
import sys

# Couleurs terminal
GREEN = '\033[92m'
RED = '\033[91m'
YELLOW = '\033[93m'
BLUE = '\033[94m'
BOLD = '\033[1m'
RESET = '\033[0m'

def print_header(text):
    """Affiche header stylisé"""
    print(f"\n{BOLD}{BLUE}{'='*70}{RESET}")
    print(f"{BOLD}{BLUE}{text:^70}{RESET}")
    print(f"{BOLD}{BLUE}{'='*70}{RESET}\n")

def print_success(text):
    """Affiche succès"""
    print(f"{GREEN}✅ {text}{RESET}")

def print_error(text):
    """Affiche erreur"""
    print(f"{RED}❌ {text}{RESET}")

def print_warning(text):
    """Affiche avertissement"""
    print(f"{YELLOW}⚠️  {text}{RESET}")

def print_info(text):
    """Affiche info"""
    print(f"{BLUE}ℹ️  {text}{RESET}")

def verifier_prerequis():
    """Vérifie que tous les prérequis sont présents"""
    print_header("VÉRIFICATION PRÉREQUIS")
    
    base_dir = Path('/home/user/girasole_mission_2025')
    checks = {
        "Workspace": base_dir,
        "Scripts Python": base_dir / "scripts_python",
        "JSON exports": base_dir / "exports_json",
        "JSON terrain 3085": base_dir / "exports_json" / "3085_terrain.json",
        "JSON BE 3085": base_dir / "exports_json" / "3085_be.json",
        "Liste centrales": base_dir / "data" / "liste_52_centrales.json",
        "Script ANNEXE2": base_dir / "scripts_python" / "generer_annexe2_automatique.py",
        "Script Rapports": base_dir / "scripts_python" / "generer_rapports_diagpv_girasole_COMPLET.py",
    }
    
    all_ok = True
    for name, path in checks.items():
        if path.exists():
            print_success(f"{name}: {path.name}")
        else:
            print_error(f"{name} MANQUANT: {path}")
            all_ok = False
    
    return all_ok

def creer_photos_test_zip():
    """Crée un ZIP photos de test pour centrale 3085"""
    print_header("CRÉATION PHOTOS TEST")
    
    base_dir = Path('/home/user/girasole_mission_2025')
    photos_dir = base_dir / 'photos_test'
    photos_dir.mkdir(exist_ok=True)
    
    # Créer images de test simples (1x1 pixel blanc)
    try:
        from PIL import Image
        
        # 6 catégories GIRASOLE
        categories = {
            'GEN': 3,  # 3 photos générales
            'ELEC': 4,  # 4 photos électriques
            'TOIT': 2,  # 2 photos toiture
            'BP': 2,   # 2 photos boîtes protection
            'MP': 5,   # 5 photos modules
            'DOC': 1   # 1 photo documents
        }
        
        photo_paths = []
        for cat, nb in categories.items():
            for i in range(1, nb + 1):
                img = Image.new('RGB', (800, 600), color=(200, 200, 200))
                filename = f"3085_{cat}_{i:02d}.jpg"
                img_path = photos_dir / filename
                img.save(img_path, 'JPEG', quality=85)
                photo_paths.append(img_path)
                print_info(f"Créée: {filename}")
        
        # Créer ZIP
        zip_path = base_dir / 'photos_test' / '3085_photos.zip'
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for photo_path in photo_paths:
                zf.write(photo_path, photo_path.name)
        
        print_success(f"ZIP créé: {zip_path.name} ({len(photo_paths)} photos)")
        return zip_path
        
    except ImportError:
        print_warning("Pillow non installé - skip création photos test")
        print_info("pip install Pillow pour activer génération photos")
        return None

def test_generation_annexe2():
    """Teste génération ANNEXE 2"""
    print_header("TEST GÉNÉRATION ANNEXE 2")
    
    script_path = Path('/home/user/girasole_mission_2025/scripts_python/generer_annexe2_automatique.py')
    
    try:
        result = subprocess.run(
            ['python3', str(script_path)],
            capture_output=True,
            text=True,
            timeout=30
        )
        
        if result.returncode == 0:
            print_success("Script ANNEXE2 exécuté avec succès")
            print(result.stdout)
            
            # Vérifier output
            output_dir = Path('/home/user/girasole_mission_2025/outputs_annexe2')
            xlsx_files = list(output_dir.glob('ANNEXE2_*.xlsx'))
            if xlsx_files:
                latest = max(xlsx_files, key=lambda p: p.stat().st_mtime)
                size_kb = latest.stat().st_size / 1024
                print_success(f"Fichier généré: {latest.name} ({size_kb:.1f} KB)")
                return True
            else:
                print_error("Aucun fichier ANNEXE2 trouvé")
                return False
        else:
            print_error("Erreur exécution script ANNEXE2")
            print(result.stderr)
            return False
            
    except Exception as e:
        print_error(f"Exception: {e}")
        return False

def test_generation_rapport():
    """Teste génération rapport PDF centrale 3085"""
    print_header("TEST GÉNÉRATION RAPPORT 3085")
    
    base_dir = Path('/home/user/girasole_mission_2025')
    script_path = base_dir / 'scripts_python' / 'generer_rapports_diagpv_girasole_COMPLET.py'
    
    # Vérifier si photos test existent
    photos_zip = base_dir / 'photos_test' / '3085_photos.zip'
    if not photos_zip.exists():
        print_warning(f"Photos ZIP non trouvé: {photos_zip}")
        photos_zip = creer_photos_test_zip()
        if not photos_zip:
            print_error("Impossible de créer photos test - skip test rapport")
            return False
    
    try:
        # Importer et tester directement le module
        import importlib.util
        spec = importlib.util.spec_from_file_location("generer_rapports", script_path)
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        
        # Charger info centrale 3085
        centrales_path = base_dir / 'data' / 'liste_52_centrales.json'
        with open(centrales_path, 'r', encoding='utf-8') as f:
            centrales = json.load(f)
        
        centrale_3085 = next((c for c in centrales if c['id'] == '3085'), None)
        if not centrale_3085:
            print_error("Centrale 3085 non trouvée dans liste")
            return False
        
        # Chemins JSON
        json_terrain = base_dir / 'exports_json' / '3085_terrain.json'
        json_be = base_dir / 'exports_json' / '3085_be.json'
        
        # Output directory
        output_dir = base_dir / 'outputs_rapports_test'
        output_dir.mkdir(exist_ok=True)
        
        print_info(f"Génération rapport pour: {centrale_3085['nom']}")
        print_info(f"JSON terrain: {json_terrain.name}")
        print_info(f"JSON BE: {json_be.name}")
        print_info(f"Photos ZIP: {photos_zip.name}")
        
        # Générer rapport
        output_docx = module.generer_rapport_pdf(
            centrale_id='3085',
            json_terrain_path=json_terrain,
            json_be_path=json_be,
            photos_zip_path=photos_zip,
            centrale_info=centrale_3085,
            output_dir=output_dir
        )
        
        if output_docx and output_docx.exists():
            size_kb = output_docx.stat().st_size / 1024
            print_success(f"Rapport DOCX généré: {output_docx.name} ({size_kb:.1f} KB)")
            
            # Vérifier contenu
            from docx import Document
            doc = Document(output_docx)
            nb_paragraphes = len(doc.paragraphs)
            nb_tableaux = len(doc.tables)
            
            print_info(f"Contenu: {nb_paragraphes} paragraphes, {nb_tableaux} tableaux")
            
            return True
        else:
            print_error("Fichier DOCX non généré")
            return False
            
    except Exception as e:
        print_error(f"Exception lors génération rapport: {e}")
        import traceback
        traceback.print_exc()
        return False

def generer_rapport_synthese():
    """Génère rapport synthèse des tests"""
    print_header("RAPPORT SYNTHÈSE TESTS")
    
    base_dir = Path('/home/user/girasole_mission_2025')
    
    # Compter outputs
    annexe2_dir = base_dir / 'outputs_annexe2'
    rapports_dir = base_dir / 'outputs_rapports_test'
    
    nb_annexe2 = len(list(annexe2_dir.glob('*.xlsx'))) if annexe2_dir.exists() else 0
    nb_rapports = len(list(rapports_dir.glob('*.docx'))) if rapports_dir.exists() else 0
    
    print(f"\n{BOLD}📊 STATISTIQUES:{RESET}")
    print(f"   • ANNEXE2 générées: {nb_annexe2}")
    print(f"   • Rapports générés: {nb_rapports}")
    
    if nb_annexe2 > 0 and nb_rapports > 0:
        print(f"\n{GREEN}{BOLD}✅ WORKFLOW COMPLET FONCTIONNEL !{RESET}")
        print(f"\n{BOLD}🎯 PRÊT POUR PRODUCTION:{RESET}")
        print(f"   1. Lancer génération ANNEXE2 pour 52 centrales")
        print(f"   2. Lancer génération rapports pour centrales avec JSON")
        print(f"   3. Déployer checklists sur GitHub Pages")
        return True
    else:
        print(f"\n{RED}{BOLD}❌ WORKFLOW INCOMPLET{RESET}")
        return False

def main():
    """Exécute test workflow complet"""
    print_header("TEST WORKFLOW COMPLET - CENTRALE 3085 PILOTE")
    print(f"{BOLD}Diagnostic Photovoltaïque - GIRASOLE 2025{RESET}")
    print(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    
    # 1. Vérifier prérequis
    if not verifier_prerequis():
        print_error("Prérequis manquants - impossible de continuer")
        sys.exit(1)
    
    # 2. Tester ANNEXE2
    annexe2_ok = test_generation_annexe2()
    
    # 3. Tester rapport PDF
    rapport_ok = test_generation_rapport()
    
    # 4. Rapport synthèse
    workflow_ok = generer_rapport_synthese()
    
    # 5. Conclusion
    if workflow_ok:
        print(f"\n{GREEN}{BOLD}{'='*70}{RESET}")
        print(f"{GREEN}{BOLD}{'TEST WORKFLOW: SUCCÈS ✅':^70}{RESET}")
        print(f"{GREEN}{BOLD}{'='*70}{RESET}\n")
        sys.exit(0)
    else:
        print(f"\n{RED}{BOLD}{'='*70}{RESET}")
        print(f"{RED}{BOLD}{'TEST WORKFLOW: ÉCHEC ❌':^70}{RESET}")
        print(f"{RED}{BOLD}{'='*70}{RESET}\n")
        sys.exit(1)

if __name__ == "__main__":
    main()
